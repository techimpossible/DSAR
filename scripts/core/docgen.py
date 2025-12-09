"""
Word document generation for DSAR reports.

This module provides functions to generate:
- Vendor-specific DSAR reports (create_vendor_report)
- GDPR-compliant cover letters (create_cover_letter)
- Internal redaction keys (create_redaction_key)

All documents are generated using python-docx and follow
GDPR Article 15 requirements for data subject access requests.
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
from typing import Dict, List, Any, Optional


def create_vendor_report(
    vendor_name: str,
    data_subject_name: str,
    data_subject_email: str,
    profile_data: Dict[str, Any],
    records: List[Dict],
    categories: List[str] = None,
    redaction_stats: Dict[str, int] = None,
    export_filename: str = None
) -> Document:
    """
    Generate a standardized DSAR report for any vendor.

    Args:
        vendor_name: Name of the vendor system (e.g., "Slack", "HubSpot")
        data_subject_name: Full name of the data subject
        data_subject_email: Email address of the data subject
        profile_data: Dictionary of profile/account information
        records: List of activity records (each with date, type, category, content)
        categories: List of memberships/categories (e.g., Slack channels)
        redaction_stats: Statistics from RedactionEngine.get_stats()
        export_filename: Name of the source export file

    Returns:
        A python-docx Document object ready to be saved
    """
    doc = Document()

    # Title
    title = doc.add_heading(f'Data Subject Access Request - {vendor_name}', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph()

    # Summary section
    doc.add_heading('Summary', level=1)

    summary_table = doc.add_table(rows=5, cols=2)
    summary_table.style = 'Table Grid'

    summary_data = [
        ('Data Subject', data_subject_name),
        ('Email', data_subject_email or 'N/A'),
        ('Records Found', str(len(records))),
        ('Export Source', export_filename or 'N/A'),
        ('Report Generated', datetime.now().strftime('%Y-%m-%d')),
    ]

    for i, (label, value) in enumerate(summary_data):
        summary_table.rows[i].cells[0].text = label
        summary_table.rows[i].cells[1].text = value

    doc.add_paragraph()

    # Profile data section
    if profile_data:
        doc.add_heading('Profile Data', level=1)

        profile_table = doc.add_table(rows=len(profile_data), cols=2)
        profile_table.style = 'Table Grid'

        for i, (key, value) in enumerate(profile_data.items()):
            profile_table.rows[i].cells[0].text = str(key)
            profile_table.rows[i].cells[1].text = str(value) if value else 'N/A'

        doc.add_paragraph()

    # Categories/memberships section
    if categories:
        doc.add_heading('Memberships / Categories', level=1)
        for cat in categories:
            doc.add_paragraph(f"• {cat}", style='List Bullet')
        doc.add_paragraph()

    # Records section
    doc.add_heading('Activity Records', level=1)

    if records:
        # Create table with headers
        headers = ['Date', 'Type', 'Category', 'Content']
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = 'Table Grid'

        # Header row
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            # Make header bold
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True

        # Data rows (limit to prevent massive docs)
        max_records = 500
        for record in records[:max_records]:
            row = table.add_row()
            row.cells[0].text = _truncate(str(record.get('date', 'N/A')), 20)
            row.cells[1].text = _truncate(str(record.get('type', 'N/A')), 30)
            row.cells[2].text = _truncate(str(record.get('category', 'N/A')), 50)
            row.cells[3].text = _truncate(str(record.get('content', 'N/A')), 500)

        if len(records) > max_records:
            doc.add_paragraph()
            doc.add_paragraph(
                f"Note: Showing {max_records} of {len(records)} records. "
                f"See JSON export for complete data."
            )
    else:
        doc.add_paragraph("No activity records found.")

    doc.add_paragraph()

    # Redaction note
    doc.add_heading('Redaction Note', level=1)
    doc.add_paragraph(
        "In accordance with GDPR Article 15(4), personal data relating to other individuals "
        "has been redacted from this report. Redacted information is indicated by placeholders "
        "such as [REDACTED_USER_1], [REDACTED_EMAIL_1], etc. This ensures the privacy rights "
        "of third parties are protected while providing complete access to your personal data."
    )

    if redaction_stats:
        doc.add_paragraph()
        doc.add_paragraph("Redaction summary:")
        for category, count in redaction_stats.items():
            if count > 0:
                doc.add_paragraph(
                    f"• {category.title()}s redacted: {count}",
                    style='List Bullet'
                )

    return doc


def create_cover_letter(
    data_subject_name: str,
    data_subject_email: str,
    vendors_processed: Dict[str, int],
    request_date: str,
    company_name: str,
    dpo_name: str = "Data Protection Officer",
    dpo_email: str = None,
    company_address: str = None
) -> Document:
    """
    Generate a GDPR-compliant cover letter for the DSAR response package.

    Args:
        data_subject_name: Full name of the data subject
        data_subject_email: Email address of the data subject
        vendors_processed: Dictionary mapping vendor names to record counts
        request_date: Date the original DSAR was received (e.g., "15 January 2025")
        company_name: Name of the responding organization
        dpo_name: Name of the Data Protection Officer
        dpo_email: Email of the DPO
        company_address: Physical address of the organization

    Returns:
        A python-docx Document object ready to be saved
    """
    doc = Document()

    # Header
    doc.add_heading('DATA SUBJECT ACCESS REQUEST RESPONSE', 0)

    # Generate reference number from initials
    name_parts = data_subject_name.split()
    if len(name_parts) >= 2:
        ref_initials = f"{name_parts[0][:2]}{name_parts[-1][:2]}".upper()
    else:
        ref_initials = name_parts[0][:4].upper() if name_parts else "DSAR"

    doc.add_paragraph(f"Reference: DSAR-{datetime.now().strftime('%Y%m%d')}-{ref_initials}")
    doc.add_paragraph(f"Date: {datetime.now().strftime('%d %B %Y')}")
    doc.add_paragraph()

    # Salutation
    doc.add_paragraph(f"Dear {data_subject_name},")
    doc.add_paragraph()
    doc.add_paragraph(f"RE: Your Data Subject Access Request dated {request_date}")
    doc.add_paragraph()

    # Introduction
    doc.add_paragraph(
        "We write in response to your request for access to your personal data pursuant to "
        "Article 15 of the General Data Protection Regulation (GDPR)."
    )
    doc.add_paragraph()

    # Summary
    doc.add_heading('Summary of Search Conducted', level=1)

    total_records = sum(vendors_processed.values())
    doc.add_paragraph(
        f"We have searched {len(vendors_processed)} systems and identified {total_records} "
        f"records containing your personal data:"
    )
    doc.add_paragraph()

    for vendor, count in sorted(vendors_processed.items()):
        doc.add_paragraph(f"• {vendor}: {count} records", style='List Bullet')

    doc.add_paragraph()

    # Enclosed documents
    doc.add_heading('Enclosed Documents', level=1)
    doc.add_paragraph("Please find enclosed:")
    doc.add_paragraph(
        "1. Individual reports for each system searched, containing your profile data "
        "and activity records",
        style='List Number'
    )
    doc.add_paragraph(
        "2. Raw data exports in machine-readable format (JSON)",
        style='List Number'
    )
    doc.add_paragraph()

    # Redaction explanation
    doc.add_heading('Redaction of Third-Party Data', level=1)
    doc.add_paragraph(
        "In accordance with Article 15(4) of the GDPR, which states that the right to obtain "
        "a copy of personal data 'shall not adversely affect the rights and freedoms of others,' "
        "we have redacted personal data relating to other individuals."
    )
    doc.add_paragraph()
    doc.add_paragraph(
        "Redacted information is indicated by placeholders such as [REDACTED_USER_1]. This ensures "
        "we protect the privacy rights of other data subjects while providing you with complete "
        "access to your own personal data."
    )
    doc.add_paragraph()
    doc.add_paragraph(
        "Note: Redaction labels are vendor-specific. The same individual may have different "
        "labels in different system reports."
    )
    doc.add_paragraph()

    # Rights reminder
    doc.add_heading('Your Rights', level=1)
    doc.add_paragraph("Under the GDPR, you also have the following rights:")

    rights = [
        "Right to rectification (Article 16) - to have inaccurate data corrected",
        "Right to erasure (Article 17) - to have your data deleted in certain circumstances",
        "Right to restriction (Article 18) - to limit how we use your data",
        "Right to data portability (Article 20) - to receive your data in a portable format",
        "Right to object (Article 21) - to object to certain processing activities",
        "Right to lodge a complaint with a supervisory authority",
    ]

    for right in rights:
        doc.add_paragraph(f"• {right}", style='List Bullet')

    doc.add_paragraph()

    # Contact
    doc.add_heading('Contact', level=1)
    doc.add_paragraph("If you have questions or wish to exercise your rights, please contact:")
    doc.add_paragraph()
    doc.add_paragraph(dpo_name)
    if dpo_email:
        doc.add_paragraph(f"Email: {dpo_email}")
    if company_address:
        doc.add_paragraph(f"Address: {company_address}")

    doc.add_paragraph()
    doc.add_paragraph("Yours sincerely,")
    doc.add_paragraph()
    doc.add_paragraph(company_name)

    return doc


def create_redaction_key(
    redaction_map: Dict[str, str],
    data_subject_name: str
) -> Document:
    """
    Generate an internal redaction key document.

    WARNING: This document contains the actual identities of redacted individuals.
    It must NEVER be sent to the data subject and should be retained only for
    internal audit purposes.

    Args:
        redaction_map: Dictionary mapping redaction labels to original identities
        data_subject_name: Name of the data subject (for reference)

    Returns:
        A python-docx Document object ready to be saved
    """
    doc = Document()

    # Warning header
    warning = doc.add_heading(
        'INTERNAL DOCUMENT - DO NOT SEND TO DATA SUBJECT',
        0
    )
    warning.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    doc.add_heading('Redaction Key', level=1)
    doc.add_paragraph(f"DSAR for: {data_subject_name}")
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph()

    doc.add_paragraph(
        "This document maps redaction placeholders to actual identities. "
        "Retain for audit purposes. DO NOT include in DSAR response."
    )
    doc.add_paragraph()

    # Redaction table
    if redaction_map:
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'

        # Header row
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Redaction Label'
        header_cells[1].text = 'Actual Identity'
        for cell in header_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True

        # Data rows
        for label, identity in sorted(redaction_map.items()):
            row = table.add_row()
            row.cells[0].text = label
            row.cells[1].text = identity
    else:
        doc.add_paragraph("No redactions were applied.")

    return doc


def _truncate(text: str, max_length: int) -> str:
    """Truncate text to max length with ellipsis."""
    if not text:
        return ''
    text = str(text)
    if len(text) <= max_length:
        return text
    return text[:max_length - 3] + '...'
